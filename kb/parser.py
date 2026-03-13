# kb/parser.py
import os
from PyPDF2 import PdfReader
from docx import Document
from config import Config

class DocumentParser:
    """解析上传的文档，提取文本内容"""

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """解析PDF文件"""
        text = ""
        try:
            reader = PdfReader(file_path)
            print(f"📄 PDF page num: {len(reader.pages)}")
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                    print(f"   第{i+1}页: {len(page_text)}字符")
        except Exception as e:
            raise Exception(f"PDF解析失败: {e}")
        return text

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """解析DOCX文件"""
        text = ""
        try:
            doc = Document(file_path)
            print(f"📄 DOCX段落数: {len(doc.paragraphs)}")
            # 提取段落
            para_count = 0
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
                    para_count += 1
            print(f"   提取了{para_count}个段落")
            
            # 提取表格
            table_count = 0
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                        table_count += 1
            if table_count > 0:
                print(f"   提取了{table_count}个表格行")
        
        except Exception as e:
            raise Exception(f"DOCX解析失败: {e}")
        return text

    @classmethod
    def parse(self, file_path: str) -> dict:
        """统一解析接口，返回文本内容和元数据"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
            
        ext = os.path.splitext(file_path)[1].lower()
        filename = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)
        
        print(f"\n📂 开始解析文件: {filename}")
        print(f"   文件大小: {file_size} 字节")
        print(f"   文件格式: {ext}")

        if ext == '.pdf':
            text = self.parse_pdf(file_path)
        elif ext == '.docx':
            text = self.parse_docx(file_path)
        elif ext == '.txt':
            # 尝试多种编码打开文件
            encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030']
            text = None
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    print(f"   使用编码 {encoding} 成功")
                    break
                except UnicodeDecodeError:
                    continue
            if text is None:
                # 最后尝试用二进制读取并忽略错误
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                print("   使用utf-8忽略错误模式")
            print(f"📄 TXT文件: {len(text)}字符")
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

        print(f"✅ 解析完成，共提取 {len(text)} 字符\n")
        
        return {
            "text": text,
            "filename": filename,
            "file_size": file_size,
            "format": ext
        }